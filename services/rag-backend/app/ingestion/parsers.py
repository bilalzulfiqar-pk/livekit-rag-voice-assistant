from __future__ import annotations

import io
import re
from collections import Counter
from collections.abc import Iterable
from dataclasses import dataclass, field
from typing import Any

import pdfplumber
from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


TEXT_FILE_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".log",
    ".json",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".html",
    ".css",
}

PDF_TEXT_MIN_LENGTH = 40
HEADER_FOOTER_REPEAT_MIN_COUNT = 2
HEADER_FOOTER_LINE_LIMIT = 2
LINE_TOP_TOLERANCE = 4
RUNNING_HEADER_TOP_THRESHOLD = 110
RUNNING_FOOTER_BOTTOM_THRESHOLD = 760
BULLET_MARKERS = ("-", "*", "•", "·", "â€¢")


class DocumentParseError(ValueError):
    """Raised when uploaded content cannot be parsed into usable structured text."""


@dataclass(slots=True)
class ParsedBlock:
    kind: str
    text: str
    heading_path: list[str] = field(default_factory=list)
    page_number: int | None = None
    table_id: str | None = None
    row_index: int | None = None


@dataclass(slots=True)
class ParsedDocument:
    source_format: str
    parser_name: str
    normalized_text: str
    blocks: list[ParsedBlock]
    parse_metadata: dict[str, Any] = field(default_factory=dict)


def parse_uploaded_document(content: bytes, filename: str) -> ParsedDocument:
    extension = _normalized_extension(filename)

    if extension == ".pdf":
        return _parse_pdf_document(content)

    if extension == ".docx":
        return _parse_docx_document(content)

    if extension in TEXT_FILE_EXTENSIONS:
        return _parse_text_document(content, extension)

    raise DocumentParseError("Unsupported file type. Upload a UTF-8 text file, PDF, or DOCX document.")


def _parse_text_document(content: bytes, extension: str) -> ParsedDocument:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("Only UTF-8 text files, machine-generated PDFs, and DOCX files are supported.") from exc

    normalized_lines = [_normalize_inline_whitespace(line) for line in text.replace("\r\n", "\n").split("\n")]
    blocks: list[ParsedBlock] = []
    paragraph_lines: list[str] = []
    current_heading_path: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        paragraph_text = " ".join(paragraph_lines).strip()
        if paragraph_text:
            blocks.append(
                ParsedBlock(
                    kind="paragraph",
                    text=paragraph_text,
                    heading_path=list(current_heading_path),
                )
            )
        paragraph_lines.clear()

    for line in normalized_lines:
        if not line:
            flush_paragraph()
            continue

        if _looks_like_heading(line):
            flush_paragraph()
            current_heading_path = [line]
            blocks.append(ParsedBlock(kind="heading", text=line, heading_path=list(current_heading_path)))
            continue

        if _looks_like_list_item(line):
            flush_paragraph()
            blocks.append(
                ParsedBlock(
                    kind="list_item",
                    text=_normalize_list_item(line),
                    heading_path=list(current_heading_path),
                )
            )
            continue

        paragraph_lines.append(line)

    flush_paragraph()
    normalized_text = _join_blocks(blocks)
    if not normalized_text:
        raise DocumentParseError("Uploaded text file does not contain any usable text.")

    return ParsedDocument(
        source_format=extension.lstrip("."),
        parser_name="plain_text",
        normalized_text=normalized_text,
        blocks=blocks,
        parse_metadata={"parser_mode": "plain_text"},
    )


def _parse_pdf_document(content: bytes) -> ParsedDocument:
    try:
        pdf = pdfplumber.open(io.BytesIO(content))
    except Exception as exc:  # pragma: no cover - library-specific exception types vary
        raise DocumentParseError(
            "Unable to read this PDF file. Make sure it is not corrupted, encrypted, or password protected."
        ) from exc

    with pdf:
        page_records = [_extract_pdf_page_record(page, page_number=index + 1) for index, page in enumerate(pdf.pages)]

    repeated_lines = _detect_repeated_header_footer_lines(page_records)

    blocks: list[ParsedBlock] = []
    current_heading_path: list[str] = []

    for record in page_records:
        filtered_lines = [
            line["text"]
            for line in record["text_lines"]
            if line["normalized"] not in repeated_lines
            and not _is_pdf_running_header_or_footer(line, page_number=record["page_number"])
        ]

        paragraph_lines: list[str] = []

        def flush_paragraph() -> None:
            if not paragraph_lines:
                return
            paragraph_text = " ".join(paragraph_lines).strip()
            if paragraph_text:
                blocks.append(
                    ParsedBlock(
                        kind="paragraph",
                        text=paragraph_text,
                        heading_path=list(current_heading_path),
                        page_number=record["page_number"],
                    )
                )
            paragraph_lines.clear()

        for raw_line in filtered_lines:
            for line in _split_inline_bullet_segments(raw_line):
                if not line:
                    flush_paragraph()
                    continue

                if _looks_like_heading(line):
                    flush_paragraph()
                    current_heading_path = [line]
                    blocks.append(
                        ParsedBlock(
                            kind="heading",
                            text=line,
                            heading_path=list(current_heading_path),
                            page_number=record["page_number"],
                        )
                    )
                    continue

                if _looks_like_list_item(line):
                    flush_paragraph()
                    blocks.append(
                        ParsedBlock(
                            kind="list_item",
                            text=_normalize_list_item(line),
                            heading_path=list(current_heading_path),
                            page_number=record["page_number"],
                        )
                    )
                    continue

                paragraph_lines.append(line)

        flush_paragraph()

        for table_index, row_texts in enumerate(record["tables"], start=1):
            table_id = f"p{record['page_number']}t{table_index}"
            for row_index, row_text in enumerate(row_texts):
                blocks.append(
                    ParsedBlock(
                        kind="table_row",
                        text=row_text,
                        heading_path=list(current_heading_path),
                        page_number=record["page_number"],
                        table_id=table_id,
                        row_index=row_index,
                    )
                )

    blocks = _normalize_pdf_blocks(blocks)

    normalized_text = _join_blocks(blocks)
    if len(normalized_text) < PDF_TEXT_MIN_LENGTH:
        raise DocumentParseError(
            "This PDF does not contain enough machine-readable text to ingest. OCR-based PDFs are not supported."
        )

    return ParsedDocument(
        source_format="pdf",
        parser_name="pdfplumber",
        normalized_text=normalized_text,
        blocks=blocks,
        parse_metadata={
            "parser_mode": "machine_text_no_ocr",
            "page_count": len(page_records),
        },
    )


def _parse_docx_document(content: bytes) -> ParsedDocument:
    try:
        document = DocxDocument(io.BytesIO(content))
    except Exception as exc:  # pragma: no cover - library-specific exception types vary
        raise DocumentParseError(
            "Unable to read this DOCX file. Make sure it is not corrupted and uses the standard DOCX format."
        ) from exc

    blocks: list[ParsedBlock] = []
    heading_path: list[str] = []
    table_counter = 0

    for block in _iter_docx_block_items(document):
        if isinstance(block, Paragraph):
            text = _normalize_inline_whitespace(block.text)
            if not text:
                continue

            heading_level = _docx_heading_level(block)
            if heading_level is not None:
                heading_path = heading_path[: heading_level - 1]
                heading_path.append(text)
                blocks.append(
                    ParsedBlock(
                        kind="heading",
                        text=text,
                        heading_path=list(heading_path),
                    )
                )
                continue

            if _docx_is_list_item(block):
                blocks.append(
                    ParsedBlock(
                        kind="list_item",
                        text=_normalize_list_item(text),
                        heading_path=list(heading_path),
                    )
                )
                continue

            blocks.append(
                ParsedBlock(
                    kind="paragraph",
                    text=text,
                    heading_path=list(heading_path),
                )
            )
            continue

        if isinstance(block, Table):
            table_counter += 1
            table_rows = _serialize_table_rows(
                [[_normalize_inline_whitespace(cell.text) for cell in row.cells] for row in block.rows]
            )
            for row_index, row_text in enumerate(table_rows):
                blocks.append(
                    ParsedBlock(
                        kind="table_row",
                        text=row_text,
                        heading_path=list(heading_path),
                        table_id=f"t{table_counter}",
                        row_index=row_index,
                    )
                )

    normalized_text = _join_blocks(blocks)
    if not normalized_text:
        raise DocumentParseError("This DOCX file does not contain any usable text.")

    return ParsedDocument(
        source_format="docx",
        parser_name="python-docx",
        normalized_text=normalized_text,
        blocks=blocks,
        parse_metadata={"parser_mode": "native_structure"},
    )


def _extract_pdf_page_record(page, *, page_number: int) -> dict[str, Any]:
    table_objects = page.find_tables()
    table_bboxes = [table.bbox for table in table_objects]
    text_lines = _extract_pdf_lines_outside_tables(page, table_bboxes)
    tables = [_serialize_table_rows(table.extract() or []) for table in table_objects]

    return {
        "page_number": page_number,
        "text_lines": text_lines,
        "tables": [rows for rows in tables if rows],
    }


def _extract_pdf_lines_outside_tables(page, table_bboxes: list[tuple[float, float, float, float]]) -> list[dict[str, Any]]:
    words = page.extract_words(use_text_flow=True, keep_blank_chars=False) or []
    filtered_words = [
        word
        for word in words
        if not _word_inside_any_bbox(word, table_bboxes)
    ]

    if not filtered_words:
        fallback_text = page.extract_text() or ""
        return [
            {"text": normalized, "normalized": _normalize_for_repetition(normalized)}
            for normalized in (
                _normalize_inline_whitespace(line)
                for line in fallback_text.splitlines()
            )
            if normalized
        ]

    grouped_lines: list[dict[str, Any]] = []
    current_line: list[dict[str, Any]] = []
    current_top: float | None = None

    for word in sorted(filtered_words, key=lambda item: (round(item["top"], 1), item["x0"])):
        word_top = float(word["top"])
        if current_line and current_top is not None and abs(word_top - current_top) > LINE_TOP_TOLERANCE:
            grouped_lines.append(_build_pdf_line(current_line))
            current_line = []
            current_top = None

        current_line.append(word)
        current_top = word_top if current_top is None else current_top

    if current_line:
        grouped_lines.append(_build_pdf_line(current_line))

    return grouped_lines


def _build_pdf_line(words: list[dict[str, Any]]) -> dict[str, Any]:
    text = _normalize_inline_whitespace(" ".join(word["text"] for word in sorted(words, key=lambda item: item["x0"])))
    return {
        "text": text,
        "normalized": _normalize_for_repetition(text),
        "top": min(float(word["top"]) for word in words),
        "bottom": max(float(word["bottom"]) for word in words),
    }


def _detect_repeated_header_footer_lines(page_records: list[dict[str, Any]]) -> set[str]:
    line_counter: Counter[str] = Counter()

    for record in page_records:
        page_lines = [line["normalized"] for line in record["text_lines"] if line["normalized"]]
        candidates = set(page_lines[:HEADER_FOOTER_LINE_LIMIT] + page_lines[-HEADER_FOOTER_LINE_LIMIT:])
        line_counter.update(candidate for candidate in candidates if candidate)

    return {
        line
        for line, count in line_counter.items()
        if count >= HEADER_FOOTER_REPEAT_MIN_COUNT
    }


def _normalize_pdf_blocks(blocks: list[ParsedBlock]) -> list[ParsedBlock]:
    normalized_blocks: list[ParsedBlock] = []

    for block in blocks:
        if normalized_blocks and block.kind == "paragraph":
            previous = normalized_blocks[-1]
            if (
                previous.kind == "list_item"
                and previous.heading_path == block.heading_path
                and previous.page_number == block.page_number
                and _looks_like_list_continuation(block.text)
            ):
                previous.text = f"{previous.text.rstrip()} {block.text.lstrip()}".strip()
                continue

        normalized_blocks.append(block)

    return normalized_blocks


def _looks_like_list_continuation(text: str) -> bool:
    cleaned = _normalize_inline_whitespace(text)
    if not cleaned:
        return False
    lowered = cleaned.lower()
    return lowered[:1].islower() or lowered.startswith(
        ("aren't", "aren’t", "isn't", "isn’t", "doesn't", "doesn’t", "not")
    )


def _is_pdf_running_header_or_footer(line: dict[str, Any], *, page_number: int) -> bool:
    text = line["text"]
    normalized = line["normalized"]
    top = float(line.get("top", 0.0))
    bottom = float(line.get("bottom", 0.0))

    if re.search(r"\.{4,}\s*\d+\s*$", text):
        return True

    if top <= RUNNING_HEADER_TOP_THRESHOLD:
        if re.match(r"^(?:chapter|section)\b.*\s+\d+\s*$", normalized):
            return True
        if normalized in {"myaarpmedicare.com", "questions??"}:
            return True
        if page_number > 1 and re.match(r"^(?:chapter|section)\b", normalized) and len(normalized) <= 80:
            return True

    if bottom >= RUNNING_FOOTER_BOTTOM_THRESHOLD:
        if re.fullmatch(r"\d+", normalized):
            return True
        if "questions??" in normalized:
            return True

    return False


def _split_inline_bullet_segments(text: str) -> list[str]:
    cleaned = _normalize_inline_whitespace(text)
    if not cleaned:
        return []

    split_text = re.sub(r"\s+(?=[•·])", "\n", cleaned)
    return [segment.strip() for segment in split_text.split("\n") if segment.strip()]


def _serialize_table_rows(rows: Iterable[Iterable[str | None]]) -> list[str]:
    normalized_rows = [
        [_normalize_inline_whitespace(cell or "") for cell in row]
        for row in rows
    ]
    normalized_rows = [[cell for cell in row if cell] for row in normalized_rows]
    normalized_rows = [row for row in normalized_rows if row]
    if not normalized_rows:
        return []

    if len(normalized_rows[0]) >= 3 and _looks_like_header_row(normalized_rows[0]):
        headers = normalized_rows[0]
        data_rows = normalized_rows[1:]
        serialized: list[str] = []
        for row in data_rows:
            pairs = [
                f"{header}: {value}"
                for header, value in zip(headers, row)
                if header and value
            ]
            if pairs:
                serialized.append(" | ".join(pairs))
        if serialized:
            return serialized

    serialized_rows: list[str] = []
    for row in normalized_rows:
        if len(row) == 1:
            serialized_rows.append(row[0])
        elif len(row) == 2:
            serialized_rows.append(f"{row[0]}: {row[1]}")
        else:
            serialized_rows.append(" | ".join(row))
    return serialized_rows


def _looks_like_header_row(row: list[str]) -> bool:
    if not row:
        return False
    alpha_heavy = sum(1 for cell in row if re.search(r"[A-Za-z]", cell) and not re.search(r"\b\d+\b", cell))
    return alpha_heavy >= max(2, len(row) - 1)


def _iter_docx_block_items(document: DocxDocumentType) -> Iterable[Paragraph | Table]:
    parent = document.element.body
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _docx_heading_level(paragraph: Paragraph) -> int | None:
    style_name = _normalize_inline_whitespace(getattr(paragraph.style, "name", "") or "").lower()
    match = re.match(r"heading\s+(\d+)", style_name)
    if not match:
        return None
    return int(match.group(1))


def _docx_is_list_item(paragraph: Paragraph) -> bool:
    style_name = _normalize_inline_whitespace(getattr(paragraph.style, "name", "") or "").lower()
    return "list" in style_name or _looks_like_list_item(paragraph.text)


def _join_blocks(blocks: list[ParsedBlock]) -> str:
    return "\n\n".join(block.text.strip() for block in blocks if block.text.strip()).strip()


def _normalized_extension(filename: str) -> str:
    filename = filename.strip().lower()
    if "." not in filename:
        return ""
    return filename[filename.rfind(".") :]


def _normalize_inline_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _normalize_for_repetition(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip().lower()


def _normalize_list_item(text: str) -> str:
    stripped = re.sub(r"^(?:[-*•·â€¢]\s*|\d+[.)]\s*)", "", text).strip()
    return f"- {stripped}" if stripped else ""


def _looks_like_list_item(text: str) -> bool:
    return bool(re.match(r"^(?:[-*•·â€¢]\s+|\d+[.)]\s+)", text))


def _looks_like_heading(text: str) -> bool:
    normalized = _normalize_inline_whitespace(text)
    if not normalized or len(normalized) > 120:
        return False
    if normalized.startswith(BULLET_MARKERS):
        return False
    if re.search(r"\.{4,}\s*\d+\s*$", normalized):
        return False
    if re.search(r"\b(?:chapter|section)\b.*\s+\d+\s*$", normalized, re.IGNORECASE):
        return False
    if re.search(r"\b(section|chapter|appendix|summary|overview)\b", normalized, re.IGNORECASE):
        return True
    if re.search(r"[.!?]$", normalized):
        return False
    if normalized.endswith(","):
        return False
    if normalized.count(",") >= 2 and ":" not in normalized:
        return False
    words = normalized.split()
    if len(words) > 12:
        return False
    uppercase_words = sum(1 for word in words if word[:1].isupper() or word[:1].isdigit())
    return uppercase_words >= max(1, len(words) - 2)


def _word_inside_any_bbox(word: dict[str, Any], bboxes: list[tuple[float, float, float, float]]) -> bool:
    for bbox in bboxes:
        if _word_inside_bbox(word, bbox):
            return True
    return False


def _word_inside_bbox(word: dict[str, Any], bbox: tuple[float, float, float, float]) -> bool:
    x0, top, x1, bottom = bbox
    return (
        float(word["x0"]) >= x0
        and float(word["x1"]) <= x1
        and float(word["top"]) >= top
        and float(word["bottom"]) <= bottom
    )
