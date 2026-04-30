import io
import unittest
from datetime import datetime, timezone
from types import SimpleNamespace

from docx import Document as DocxDocument
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.pdfgen import canvas

from app.documents.service import DocumentManagementService
from app.ingestion.parsers import DocumentParseError, parse_uploaded_document


def _build_two_page_pdf() -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)

    page_bodies = [
        "Primary care office visits are covered when medically necessary.",
        "Specialist office visits may require cost sharing under the plan.",
    ]

    for body in page_bodies:
        pdf.drawString(72, 760, "AARP Medicare Advantage")
        pdf.drawString(72, 720, body)
        pdf.drawString(72, 36, "Page footer 2026")
        pdf.showPage()

    pdf.save()
    return buffer.getvalue()


def _build_table_pdf() -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    table = Table(
        [
            ["Benefit", "In-network", "Out-of-network"],
            ["Primary care office visits", "$0 copayment per visit", "$25 copayment per visit"],
            ["Specialist office visits", "$40 copayment per visit", "$70 copayment per visit"],
        ]
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ]
        )
    )
    doc.build([table, Spacer(1, 12)])
    return buffer.getvalue()


def _build_blank_pdf() -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    pdf.showPage()
    pdf.save()
    return buffer.getvalue()


def _build_docx_file() -> bytes:
    document = DocxDocument()
    document.add_heading("Benefits and Costs", level=1)
    document.add_paragraph("Primary care visits are covered under the plan.")
    document.add_paragraph("Review the provider directory before you travel.", style="List Bullet")

    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Monthly plan premium"
    table.rows[0].cells[1].text = "$32.00"
    table.rows[1].cells[0].text = "Maximum out-of-pocket"
    table.rows[1].cells[1].text = "$3,900"

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class DocumentParserTests(unittest.TestCase):
    def test_pdf_parser_removes_repeated_header_and_footer_lines(self):
        parsed = parse_uploaded_document(_build_two_page_pdf(), "benefits.pdf")

        self.assertEqual(parsed.source_format, "pdf")
        self.assertEqual(parsed.parser_name, "pdfplumber")
        self.assertIn("Primary care office visits are covered when medically necessary.", parsed.normalized_text)
        self.assertIn("Specialist office visits may require cost sharing under the plan.", parsed.normalized_text)
        self.assertNotIn("AARP Medicare Advantage", parsed.normalized_text)
        self.assertNotIn("Page footer 2026", parsed.normalized_text)

    def test_pdf_parser_serializes_table_rows_semantically(self):
        parsed = parse_uploaded_document(_build_table_pdf(), "costs.pdf")

        self.assertTrue(any(block.kind == "table_row" for block in parsed.blocks))
        self.assertIn("Benefit: Primary care office visits", parsed.normalized_text)
        self.assertIn("In-network: $0 copayment per visit", parsed.normalized_text)
        self.assertIn("Out-of-network: $25 copayment per visit", parsed.normalized_text)

    def test_pdf_parser_rejects_low_text_pdf_without_ocr(self):
        with self.assertRaises(DocumentParseError) as context:
            parse_uploaded_document(_build_blank_pdf(), "scan.pdf")

        self.assertIn("machine-readable text", str(context.exception))

    def test_pdf_parser_rejects_corrupted_pdf_cleanly(self):
        with self.assertRaises(DocumentParseError) as context:
            parse_uploaded_document(b"%PDF-1.4 corrupted content", "broken.pdf")

        self.assertIn("Unable to read this PDF file", str(context.exception))

    def test_pdf_parser_ignores_running_header_and_recognizes_middle_dot_bullets(self):
        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=letter)
        pdf.drawString(72, 760, "Section 2.1 Eligibility requirements")
        pdf.drawString(72, 720, "You are eligible for membership in our plan as long as you meet all these conditions:")
        pdf.showPage()
        pdf.drawString(72, 760, "Chapter 1: Get started as a member 6")
        pdf.drawString(72, 720, "· You have both Medicare Part A and Medicare Part B")
        pdf.drawString(72, 700, "· You live in our geographic service area")
        pdf.save()

        parsed = parse_uploaded_document(buffer.getvalue(), "eligibility.pdf")

        self.assertNotIn("Chapter 1: Get started as a member 6", parsed.normalized_text)
        list_items = [block.text for block in parsed.blocks if block.kind == "list_item"]
        self.assertTrue(any("You have both Medicare Part A and Medicare Part B" in item for item in list_items))
        self.assertTrue(any("You live in our geographic service area" in item for item in list_items))

    def test_pdf_parser_keeps_wrapped_county_lines_in_same_paragraph(self):
        buffer = io.BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=letter)
        pdf.drawString(72, 760, "Section 2.2 Plan service area")
        pdf.drawString(72, 720, "Our service area includes these counties in Indiana: Bartholomew, Benton, Blackford, Boone,")
        pdf.drawString(72, 700, "Brown, Carroll, Clinton, Decatur, Delaware, Fayette, Fountain, Hamilton, Hancock, Hendricks,")
        pdf.drawString(72, 680, "Henry, Howard, Johnson, Lawrence, Madison, Marion, Monroe, Montgomery, Morgan, Orange.")
        pdf.save()

        parsed = parse_uploaded_document(buffer.getvalue(), "service-area.pdf")

        heading_texts = [block.text for block in parsed.blocks if block.kind == "heading"]
        paragraph_texts = [block.text for block in parsed.blocks if block.kind == "paragraph"]

        self.assertEqual(heading_texts, ["Section 2.2 Plan service area"])
        self.assertEqual(len(paragraph_texts), 1)
        self.assertIn("Brown, Carroll, Clinton", paragraph_texts[0])
        self.assertIn("Henry, Howard, Johnson", paragraph_texts[0])

    def test_docx_parser_preserves_document_order_and_structure(self):
        parsed = parse_uploaded_document(_build_docx_file(), "plan.docx")

        self.assertEqual(parsed.source_format, "docx")
        self.assertEqual(parsed.parser_name, "python-docx")
        self.assertEqual(parsed.blocks[0].kind, "heading")
        self.assertEqual(parsed.blocks[0].heading_path, ["Benefits and Costs"])
        self.assertEqual(parsed.blocks[1].kind, "paragraph")
        self.assertEqual(parsed.blocks[2].kind, "list_item")
        self.assertEqual(parsed.blocks[3].kind, "table_row")
        self.assertIn("Monthly plan premium: $32.00", parsed.normalized_text)

    def test_docx_parser_rejects_corrupted_docx_cleanly(self):
        with self.assertRaises(DocumentParseError) as context:
            parse_uploaded_document(b"not-a-real-docx", "broken.docx")

        self.assertIn("Unable to read this DOCX file", str(context.exception))

    def test_document_summary_reads_source_format_and_parser_name(self):
        row = SimpleNamespace(
            id=7,
            filename="benefits.pdf",
            created_at=datetime.now(timezone.utc),
            updated_at=datetime.now(timezone.utc),
            chunk_count=12,
            embedded_chunk_count=12,
        )

        summary = DocumentManagementService._build_document_summary(
            row=row,
            metadata={
                "source_type": "file",
                "source_format": "pdf",
                "parser_name": "pdfplumber",
                "embedding_provider": "local",
                "embedding_model": "model-x",
                "embedding_dimensions": 384,
            },
        )

        self.assertEqual(summary.source_format, "pdf")
        self.assertEqual(summary.parser_name, "pdfplumber")


if __name__ == "__main__":
    unittest.main()
